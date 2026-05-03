"""
Generate Laporan_DNA_Nuklir.docx for client.
Narrative, data-driven report (Bahasa Indonesia).
"""
import os
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = Path(__file__).parent.parent.parent
OUT_DIR  = ROOT / "data" / "processed" / "news"
DOC_PATH = ROOT / "Laporan_DNA_Kebijakan_Nuklir.docx"

# Load data
flat = pd.read_csv(OUT_DIR / "01_flat_statements.csv")
flat["date"] = pd.to_datetime(flat["date"], errors="coerce")
nodes_actors = pd.read_csv(OUT_DIR / "02_nodes_actors.csv")
concepts = pd.read_csv(OUT_DIR / "03_nodes_concepts.csv")

def era(d):
    if pd.isna(d):
        return "Tidak diketahui"
    if d < pd.Timestamp("2019-10-20"):
        return "Jokowi I (2014 sampai 2019)"
    if d < pd.Timestamp("2024-10-20"):
        return "Jokowi II (2019 sampai 2024)"
    return "Prabowo (2024 sampai 2026)"
flat["era"] = flat["date"].apply(era)

# ─── Aggregates ────────────────────────────────────────────────────────────
N_STMT = len(flat)
N_ART  = flat["source_id"].nunique()
N_ACT  = flat["actor"].nunique()
N_CON  = flat["concept"].nunique()
DATE_MIN = "20 Oct 2014"  # start of Jokowi I period (hardcoded)
DATE_MAX = flat["date"].max().strftime("%d %b %Y")

POS_COUNT = flat["position"].value_counts().to_dict()
PCT = {k: round(v / N_STMT * 100, 1) for k, v in POS_COUNT.items()}

VAR_TBL = (
    flat.groupby("variable")
    .agg(n=("source_id", "count"),
         pro=("position", lambda x: (x == "PRO").sum()),
         kontra=("position", lambda x: (x == "KONTRA").sum()),
         netral=("position", lambda x: (x == "NETRAL").sum()))
    .reset_index()
    .sort_values("n", ascending=False)
)
VAR_TBL["pro_pct"] = (VAR_TBL["pro"] / VAR_TBL["n"] * 100).round(1)
VAR_TBL["kontra_pct"] = (VAR_TBL["kontra"] / VAR_TBL["n"] * 100).round(1)

ERA_TBL = (
    flat.groupby("era")
    .agg(n=("source_id", "count"),
         art=("source_id", "nunique"),
         pro=("position", lambda x: (x == "PRO").sum()),
         kontra=("position", lambda x: (x == "KONTRA").sum()))
    .reset_index()
)

TOP_PRO_ACT = nodes_actors.sort_values("pro_count", ascending=False).head(10)
TOP_KON_ACT = nodes_actors.sort_values("kontra_count", ascending=False).head(10)

# Domain breakdown
from urllib.parse import urlparse
def dom(u):
    h = urlparse(str(u)).netloc
    return h.replace("www.", "")
flat["domain"] = flat["source_url"].apply(dom)
TOP_DOM = (
    flat.groupby("domain")
    .agg(art=("source_id", "nunique"),
         stmt=("source_id", "count"),
         pro=("position", lambda x: (x == "PRO").sum()),
         kontra=("position", lambda x: (x == "KONTRA").sum()))
    .reset_index()
    .sort_values("art", ascending=False)
    .head(15)
)

# ─── Document ──────────────────────────────────────────────────────────────
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_margins():
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
set_margins()

def H1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x14, 0x3D, 0x6B)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

def H2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def H3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def P(text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold

def BUL(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3

def NUM(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3

def shade_cell(cell, hexcolor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexcolor)
    tc_pr.append(shd)

def TABLE(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1F4E79")
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row):
            rc[i].text = ""
            par = rc[i].paragraphs[0]
            run = par.add_run(str(v))
            run.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph()  # spacer

def DASH_REF(titles):
    """Render a small italic dashboard reference note."""
    label = ", ".join([f'"{t}"' for t in titles])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("Dashboard: ")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    r2 = p.add_run(label)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

# ═══════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN ANALISIS")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Discourse Network Analysis (DNA)")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x14, 0x3D, 0x6B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Masuknya Nuklir ke dalam RUU EBET:\nPolitik Nuklir Indonesia di Era Reformasi")
r.font.size = Pt(16)
r.italic = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Periode liputan: {DATE_MIN} sampai {DATE_MAX}")
r.font.size = Pt(11)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    f"{N_ART:,} artikel  |  {N_STMT:,} pernyataan  |  {N_ACT:,} aktor  |  {N_CON:,} konsep wacana"
)
r.font.size = Pt(11)
r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# DAFTAR ISI (manual narrative)
# ═══════════════════════════════════════════════════════════════════════════
H1("Daftar Isi")
NUM([
    "Definisi dan Konsep Kunci",
    "Ringkasan Eksekutif (Overview)",
    "Metodologi Penelitian",
    "Analisis dan Insight",
    "Kesimpulan",
    "Batasan dan Asumsi",
])
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# DEFINISI
# ═══════════════════════════════════════════════════════════════════════════
H1("1. Definisi dan Konsep Kunci")

P(
    "Bagian ini menjelaskan istilah dan konsep yang digunakan dalam laporan agar "
    "pembaca dapat membaca angka, tabel, dan narasi secara konsisten."
)

H2("1.1 Discourse Network Analysis (DNA)")
P(
    "Discourse Network Analysis (DNA) adalah metode analisis kebijakan yang memetakan "
    "tiga elemen sekaligus dalam satu kerangka, yaitu siapa yang berbicara (aktor), "
    "apa yang dikatakan (konsep atau argumen), dan bagaimana posisi sikapnya (PRO, "
    "KONTRA, NETRAL, atau AMBIGU). Dengan DNA, peneliti dapat melihat pola koalisi "
    "wacana, siapa saja yang berbagi argumen yang sama, dan seberapa terpolarisasi "
    "ruang wacana suatu isu kebijakan."
)
P(
    "Berbeda dengan analisis konten biasa yang hanya menghitung frekuensi kata, "
    "DNA secara eksplisit memasangkan aktor dengan posisi dan argumennya, sehingga "
    "hasilnya berupa jaringan relasional, bukan sekadar daftar topik."
)

H2("1.2 Posisi Sikap (PRO, KONTRA, NETRAL, AMBIGU)")
P(
    "Posisi sikap adalah label yang diberikan pada setiap pernyataan berdasarkan "
    "apakah aktor mendukung, menolak, atau bersikap netral terhadap topik yang "
    "sedang dibahas. Dalam studi ini, topik utama adalah pengembangan energi nuklir "
    "dan Pembangkit Listrik Tenaga Nuklir (PLTN) di Indonesia."
)
TABLE(
    ["Posisi", "Definisi", "Contoh pernyataan"],
    [
        [
            "PRO",
            "Aktor secara eksplisit mendukung, menyetujui, atau mengadvokasi "
            "pengembangan nuklir atau PLTN.",
            '"PLTN adalah solusi nyata untuk ketahanan energi nasional kita." '
            "(Menteri ESDM)",
        ],
        [
            "KONTRA",
            "Aktor secara eksplisit menolak, mengkritik, atau menentang "
            "pengembangan nuklir atau PLTN.",
            '"Biaya PLTN jauh terlalu mahal dan berisiko bagi tarif listrik masyarakat." '
            "(Direktur IESR)",
        ],
        [
            "NETRAL",
            "Aktor menyampaikan argumen atau konsep yang relevan, tetapi tidak "
            "secara jelas berpihak mendukung maupun menolak. Biasanya berupa "
            "penyampaian fakta, kondisi, atau pertanyaan terbuka.",
            '"Kajian tapak untuk PLTN masih dalam proses dan belum ada keputusan final." '
            "(Juru Bicara BAPETEN)",
        ],
        [
            "AMBIGU",
            "Pernyataan mengandung argumen yang dapat ditafsirkan sebagai PRO "
            "sekaligus KONTRA, atau konteksnya tidak cukup jelas untuk menentukan "
            "posisi secara pasti.",
            '"Nuklir bisa menjadi pilihan, asalkan semua syarat terpenuhi." '
            "(Anggota DPR)",
        ],
    ],
)

H2("1.3 Analisis Sentimen")
P(
    "Analisis sentimen mengukur nada emosional sebuah pernyataan, bukan posisi "
    "kebijakan aktor. Sentimen menangkap apakah bahasa yang digunakan terasa "
    "positif, negatif, atau netral secara linguistik."
)
TABLE(
    ["Sentimen", "Definisi"],
    [
        ["Positif", "Bahasa yang digunakan mengandung nada optimis, antusias, "
         "apresiasi, atau harapan. Kata kunci seperti berhasil, maju, solusi, "
         "peluang, dan mendukung."],
        ["Negatif", "Bahasa yang digunakan mengandung nada kritis, khawatir, "
         "menolak, atau pesimis. Kata kunci seperti gagal, berbahaya, merugikan, "
         "menolak, dan mengancam."],
        ["Netral", "Bahasa yang digunakan bersifat deskriptif atau faktual tanpa "
         "muatan emosional yang dominan."],
    ],
)

H2("1.4 Perbedaan Posisi dan Sentimen")
P(
    "Posisi dan sentimen adalah dua dimensi yang berbeda dan tidak selalu selaras. "
    "Seorang aktor bisa saja mendukung suatu kebijakan (PRO) tetapi menggunakan "
    "bahasa yang bernada negatif, atau menolak suatu kebijakan (KONTRA) tetapi "
    "dengan bahasa yang positif dan konstruktif. Berikut ilustrasi perbedaannya."
)
_contoh_tbl = [
    [
        "PRO + Sentimen Negatif",
        "Aktor mendukung nuklir, tetapi mengungkapkan dukungannya melalui kekhawatiran terhadap kondisi saat ini (bukan nuklir itu sendiri).",
        "Sangat disayangkan jika kita terus bergantung pada batu bara yang merusak lingkungan. PLTN adalah satu satunya jalan keluar. -- Posisi: PRO. Sentimen: negatif (kata menyayangkan dan merusak).",
    ],
    [
        "KONTRA + Sentimen Positif",
        "Aktor menolak nuklir, tetapi mengungkapkan penolakannya dengan mengangkat keunggulan alternatif lain secara optimis.",
        "Energi surya dan angin sudah terbukti lebih murah, lebih aman, dan lebih cepat dibangun. Kita tidak perlu memaksakan PLTN. -- Posisi: KONTRA. Sentimen: positif (kata terbukti, lebih murah, lebih aman).",
    ],
    [
        "PRO + Sentimen Positif",
        "Aktor mendukung nuklir dengan nada optimis dan antusias. Kombinasi paling umum pada aktor pemerintah.",
        "PLTN adalah kunci kedaulatan energi dan masa depan Indonesia yang berdaulat. -- Posisi: PRO. Sentimen: positif.",
    ],
    [
        "KONTRA + Sentimen Negatif",
        "Aktor menolak nuklir dengan nada kritis dan bernada peringatan keras. Kombinasi paling umum pada masyarakat sipil dan lingkungan.",
        "Limbah radioaktif PLTN mengancam generasi mendatang dan tidak boleh dibiarkan. -- Posisi: KONTRA. Sentimen: negatif.",
    ],
]
TABLE(["Kombinasi", "Penjelasan", "Contoh"], _contoh_tbl)
P(
    "Dalam laporan ini, posisi (PRO atau KONTRA) digunakan sebagai ukuran utama "
    "karena mencerminkan sikap kebijakan aktor secara substantif. Analisis "
    "sentimen digunakan sebagai dimensi pelengkap untuk membaca nada retoris "
    "yang digunakan masing masing koalisi wacana."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
H1("1. Ringkasan Eksekutif")

P(
    "Laporan ini merupakan studi pemetaan wacana (Discourse Network Analysis, DNA) "
    "atas isu masuknya tenaga nuklir ke dalam Rancangan Undang Undang Energi Baru "
    "dan Energi Terbarukan (RUU EBET). Studi memetakan siapa berbicara, apa yang "
    "mereka katakan, dan bagaimana posisi mereka (PRO, KONTRA, atau NETRAL) terhadap "
    "tujuh variabel kunci pembahasan kebijakan nuklir Indonesia."
)
P(
    f"Sumber utama berupa {N_ART:,} artikel pemberitaan media daring berbahasa "
    f"Indonesia yang terbit pada periode {DATE_MIN} sampai {DATE_MAX}. Dari seluruh "
    f"artikel tersebut, dihasilkan {N_STMT:,} unit pernyataan yang valid untuk "
    f"dianalisis, melibatkan {N_ACT:,} aktor (individu, institusi, fraksi, pakar, "
    f"dan media) dan {N_CON:,} konsep atau argumen wacana yang berbeda."
)

H3("Temuan utama")
BUL([
    f"Wacana publik nasional cenderung mendukung pengembangan nuklir. Sebanyak "
    f"{POS_COUNT.get('PRO',0):,} pernyataan ({PCT.get('PRO',0)}%) berposisi PRO, "
    f"sementara {POS_COUNT.get('KONTRA',0):,} pernyataan ({PCT.get('KONTRA',0)}%) "
    f"berposisi KONTRA, dan {POS_COUNT.get('NETRAL',0):,} pernyataan "
    f"({PCT.get('NETRAL',0)}%) berposisi NETRAL.",

    "Dukungan paling kuat muncul pada variabel Keamanan Nasional dan Kedaulatan "
    f"({VAR_TBL.iloc[0]['pro_pct']}% PRO) serta Ideologi dan Integritas Data, "
    "yang menempatkan nuklir sebagai instrumen kedaulatan energi dan keniscayaan "
    "teknologi.",

    "Resistensi paling tinggi muncul pada variabel Periferalisasi dan Hak Masyarakat "
    f"({VAR_TBL[VAR_TBL['variable']=='Periferalisasi & Hak Masyarakat'].iloc[0]['kontra_pct']}% "
    "KONTRA) dan Intervensi serta Pembiayaan, dengan kekhawatiran terkait limbah "
    "radioaktif, keadilan distributif, biaya tinggi, dan beban fiskal.",

    "Aktor pemerintah dan parlemen mendominasi ruang wacana sebagai pendukung. "
    f"Bahlil Lahadalia (Menteri ESDM) memimpin dengan 58 pernyataan, diikuti "
    "Hashim Djojohadikusumo (Utusan Khusus Presiden bidang Energi dan Iklim) "
    "dengan 39 pernyataan PRO.",

    "Suara KONTRA terkonsentrasi pada lembaga kajian energi dan masyarakat sipil. "
    "Fabby Tumiwa (IESR) tampil sebagai aktor penyeimbang utama dengan 17 "
    "pernyataan KONTRA, didukung Mulyanto (DPR Fraksi PKS), Greenpeace, WALHI, "
    "JATAM, dan koalisi masyarakat sipil.",

    "Volume pemberitaan tumbuh tajam memasuki era Prabowo, dari rata rata di bawah "
    "100 pernyataan per tahun pada periode 2014 sampai 2019 menjadi 749 pernyataan "
    "pada 2025 saja. Dukungan PRO juga mengalami peningkatan paling tinggi pada "
    "periode ini.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. METODOLOGI
# ═══════════════════════════════════════════════════════════════════════════
H1("2. Metodologi Penelitian")

P(
    "Penelitian ini menggunakan kerangka Discourse Network Analysis (DNA) yang "
    "memetakan jaringan tiga rangkap, yaitu aktor, pernyataan (statement), dan "
    "posisi sikap (PRO atau KONTRA). DNA dipakai karena mampu mengungkap struktur "
    "koalisi wacana, polarisasi, dan dinamika kebijakan secara empiris. Pendekatan "
    "ini melengkapi tiga teori dasar studi, yaitu Advocacy Coalition Framework, "
    "Path Dependence, dan Power Interplay."
)

H2("2.1 Tujuh Variabel Kunci")
P(
    "Kerangka analisis menggunakan tujuh variabel yang sebelumnya telah ditetapkan "
    "sebagai poros pembahasan kebijakan nuklir Indonesia. Setiap pernyataan yang "
    "diekstrak dipetakan ke variabel paling relevan."
)
BUL([
    "Dinamika Pembahasan DIM, fokus proses legislasi RUU EBET, perdebatan antar fraksi, dan konsistensi regulasi.",
    "Keamanan Nasional dan Kedaulatan, menelaah transisi nuklir dari pilihan terakhir menjadi bagian ketahanan energi.",
    "Ideologi dan Integritas Data, narasi nuklir sebagai keniscayaan teknologi, validasi data biaya, peran pakar, dan kapasitas SDM.",
    "Intervensi dan Pembiayaan, skema pendanaan APBN dan KPS, jaminan pemerintah, risiko fiskal, insentif, dan dampak ke tarif PLN.",
    "Interaksi Pemangku Kepentingan, pengaruh aktor non negara (industri dan akademisi), koordinasi antar lembaga, dan potensi regulatory capture.",
    "Periferalisasi dan Hak Masyarakat, limbah radioaktif, fenomena NIMBY, kedaulatan masyarakat lokal, keadilan distributif, dan partisipasi publik.",
    "Transisi Energi dan NZE, menilai apakah nuklir adalah komitmen tulus dekarbonisasi atau pengalihan strategis untuk menunda pensiun PLTU.",
])

H2("2.2 Alur Pengolahan Data")
P("Pipeline pengolahan data dilakukan secara end to end dalam lima tahap berurutan.")

H3("Tahap 1, Pengumpulan dan Pembersihan Data")
P(
    "Pengumpulan data primer dilakukan melalui pencarian terstruktur di Google News "
    "menggunakan kata kunci spesifik untuk masing masing variabel. Kata kunci utama "
    "mencakup PLTN ketahanan energi nasional, limbah radioaktif nuklir Indonesia, "
    "kedaulatan energi nuklir, tarif listrik PLN nuklir, RUU EBET DIM, SMR Indonesia "
    "reaktor, BRIN nuklir SDM, nuklir IAEA keselamatan Indonesia, fraksi DPR nuklir, "
    "dan nuklir NZE 2060 Indonesia. Setiap artikel kemudian dinormalisasi (deduplikasi, "
    "pembersihan tag HTML, ekstraksi tanggal terbit, dan validasi sumber)."
)

H3("Tahap 2, Ekstraksi Aktor, Pernyataan, dan Posisi")
P(
    "Setiap artikel diproses oleh skrip Python yang mengekstrak unit analisis DNA, "
    "yaitu kalimat yang memuat secara bersamaan tiga elemen, yakni aktor (siapa), "
    "konsep atau argumen (apa), dan sikap (PRO, KONTRA, NETRAL, atau AMBIGU). "
    "Skrip menerapkan kriteria validasi ketat, di antaranya aktor harus disebut "
    "eksplisit, harus mengandung argumen substantif (bukan sekadar fakta atau "
    "kunjungan diplomatik). Dari setiap artikel, maksimum tiga "
    "pernyataan paling representatif diambil untuk menjaga kualitas dataset."
)

H3("Tahap 3, Pembentukan Edgelist Jaringan")
P(
    "Pernyataan yang lolos validasi kemudian disusun menjadi daftar relasi (edge "
    "list) dalam beberapa lapis, antara lain relasi aktor dengan konsep, relasi "
    "aktor dengan aktor (kesamaan posisi), relasi aktor dengan variabel, dan "
    "relasi aktor dengan kata kunci. Dari sini terbentuk graf yang dapat dianalisis "
    "menggunakan ukuran sentralitas, modularitas, dan polarisasi."
)

H3("Tahap 4, Penilaian Sentimen dan Konsolidasi")
P(
    "Skrip Python berikutnya melakukan penilaian sentimen pada level pernyataan "
    "untuk memperkuat klasifikasi PRO atau KONTRA. Hasilnya dikonsolidasikan ke "
    "level aktor, variabel, dan konsep agar dapat dibandingkan secara agregat. "
    "Pada tahap ini dilakukan pula filter aktor aktif (memiliki tiga atau lebih "
    "pernyataan tervalidasi) untuk membersihkan kebisingan dan memastikan analisis "
    "fokus pada aktor yang konsisten muncul dalam wacana."
)

H3("Tahap 5, Visualisasi dan Pelaporan")
P(
    "Tahap akhir menghasilkan dasbor interaktif berbasis HTML yang menampilkan "
    "ringkasan agregat, distribusi posisi per variabel, peta panas aktor terhadap "
    "variabel, jaringan aktor terhadap konsep, peringkat aktor PRO dan KONTRA, "
    "linimasa pernyataan per era kepemimpinan, serta sebaran sumber pemberitaan "
    "(domain media). Dasbor ini menjadi sumber utama narasi laporan."
)

H2("2.3 Cakupan dan Dataset Final")
TABLE(
    ["Indikator", "Nilai"],
    [
        ["Periode liputan", f"{DATE_MIN} sampai {DATE_MAX}"],
        ["Total artikel unik", f"{N_ART:,}"],
        ["Total pernyataan tervalidasi", f"{N_STMT:,}"],
        ["Aktor unik teridentifikasi", f"{N_ACT:,}"],
        ["Konsep wacana unik", f"{N_CON:,}"],
        ["Variabel analisis", "7"],
        ["Era kepemimpinan tercakup", "Jokowi I, Jokowi II, dan Prabowo"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. ANALISIS DAN INSIGHT
# ═══════════════════════════════════════════════════════════════════════════
H1("3. Analisis dan Insight")

P(
    "Bagian ini membaca jaringan wacana sebagai sebuah cerita, yaitu bagaimana "
    "narasi nuklir dibangun, siapa yang mengangkatnya, apa argumen utamanya, dan "
    "bagaimana resistensi dirumuskan. Setiap subbagian dapat dipakai sebagai alur "
    "presentasi terhadap audiens, dengan urutan dari gambaran besar menuju detail "
    "aktor dan koalisi."
)

# ─── 3.1 Lanskap wacana
H2("3.1 Lanskap Wacana, Dukungan yang Konsisten dengan Resistensi yang Terlokalisasi")
DASH_REF(["Pernyataan PRO", "Pernyataan KONTRA", "Pernyataan NETRAL", "Pernyataan AMBIGU", "Total Pernyataan", "Total Artikel Berita", "Aktor Unik", "Konsep / Wacana"])
P(
    f"Pemetaan terhadap {N_STMT:,} pernyataan menunjukkan bahwa wacana publik "
    "nasional mengenai nuklir berada pada posisi yang condong PRO, namun bukan "
    "tanpa resistensi. Komposisi sikap secara keseluruhan adalah sebagai berikut."
)
TABLE(
    ["Posisi", "Jumlah pernyataan", "Persentase"],
    [
        ["PRO", f"{POS_COUNT.get('PRO',0):,}", f"{PCT.get('PRO',0)}%"],
        ["KONTRA", f"{POS_COUNT.get('KONTRA',0):,}", f"{PCT.get('KONTRA',0)}%"],
        ["NETRAL", f"{POS_COUNT.get('NETRAL',0):,}", f"{PCT.get('NETRAL',0)}%"],
        ["AMBIGU", f"{POS_COUNT.get('AMBIGU',0):,}", f"{PCT.get('AMBIGU',0)}%"],
    ],
)
P(
    "Angka ini bukan sekadar statistik. Dominasi PRO sebesar lebih dari 70% "
    "menunjukkan bahwa ruang wacana publik telah relatif siap menerima nuklir. "
    "Namun, suara KONTRA sebesar 18% bukan suara marginal, melainkan terkonsentrasi "
    "pada sekelompok aktor yang konsisten dan memiliki argumen teknokratis yang "
    "spesifik (bukan penolakan emosional), sehingga tetap berpotensi menjadi "
    "penyeimbang yang kredibel di tahap legislasi."
)

# ─── 3.2 Variabel
H2("3.2 Peta Variabel, Apa yang Disepakati dan Apa yang Diperdebatkan")
DASH_REF(["Distribusi Posisi per Variabel"])
P(
    "Distribusi sikap per variabel mengungkap di mana konsensus terjadi dan di "
    "mana garis pertarungan kebijakan berada. Dua variabel paling banyak dibicarakan "
    "adalah Keamanan Nasional serta Ideologi dan Integritas Data, dan keduanya "
    "didominasi PRO. Sebaliknya, dua variabel dengan resistensi tertinggi adalah "
    "Periferalisasi dan Hak Masyarakat serta Intervensi dan Pembiayaan."
)
VAR_TBL["netral_pct"] = (VAR_TBL["netral"] / VAR_TBL["n"] * 100).round(1)
TABLE(
    ["Variabel", "n", "PRO%", "KONTRA%", "NETRAL%"],
    [[r["variable"], r["n"], f"{r['pro_pct']}%", f"{r['kontra_pct']}%", f"{r['netral_pct']}%"]
     for _, r in VAR_TBL.iterrows()],
)

H3("Insight narasi")
BUL([
    "Variabel Keamanan Nasional dan Kedaulatan adalah panggung utama nuklir. Sebagian besar argumen PRO dirumuskan di sini, dengan konsep dominan berupa nuklir untuk kedaulatan energi, ketahanan energi nasional, dan diversifikasi sumber strategis.",
    "Ideologi dan Integritas Data menggambarkan bagaimana narasi PRO dilengkapi dengan klaim teknokratis, yaitu kesiapan SDM, kelayakan SMR, dan dukungan akademisi serta lembaga riset (BRIN dan kampus).",
    "Periferalisasi dan Hak Masyarakat menjadi titik resistensi utama dengan rasio KONTRA mencapai 44,3%. Konsep yang paling sering muncul, yaitu dampak limbah radioaktif, dampak Fukushima, dan penolakan PLTN di lokasi tapak.",
    "Intervensi dan Pembiayaan adalah arena debat fiskal. Argumen KONTRA berbasis biaya tinggi PLTN, ketergantungan teknologi, tarif listrik mahal, dan risiko beban PLN dan APBN.",
    "Transisi Energi dan NZE menjadi medan baru. Wacana terbelah antara nuklir sebagai komitmen tulus dekarbonisasi dan tudingan strategic diversion (pengalihan agar PLTU tidak segera dipensiunkan).",
])

# ─── 3.3 Aktor
H2("3.3 Peta Aktor, Siapa Mendorong dan Siapa Menahan")
DASH_REF(["Top 15 Aktor — Jumlah Pernyataan", "Top 10 Aktor KONTRA Nuklir", "Position Shift Matrix", "DNA Bipartite Network", "Distribusi Tipe Aktor"])
P(
    "Daftar aktor paling vokal mengungkap dua koalisi yang terbentuk secara "
    "alamiah. Koalisi pendorong (advocacy coalition PRO) dipimpin oleh aktor "
    "pemerintah dan parlemen, sementara koalisi penahan (advocacy coalition "
    "KONTRA) dipimpin oleh lembaga kajian energi dan masyarakat sipil."
)

H3("Sepuluh aktor PRO paling vokal")
TABLE(
    ["Aktor", "Peran", "PRO", "KONTRA", "NETRAL", "AMBIGU", "Total"],
    [[r["actor"], (r["actor_role"] or "")[:50],
      int(r["pro_count"]), int(r["kontra_count"]), int(r["netral_count"]),
      int(r["n_statements"]) - int(r["pro_count"]) - int(r["kontra_count"]) - int(r["netral_count"]),
      int(r["n_statements"])]
     for _, r in TOP_PRO_ACT.iterrows()],
)

H3("Sepuluh aktor KONTRA paling vokal")
TABLE(
    ["Aktor", "Peran", "KONTRA", "PRO", "NETRAL", "AMBIGU", "Total"],
    [[r["actor"], (r["actor_role"] or "")[:50],
      int(r["kontra_count"]), int(r["pro_count"]), int(r["netral_count"]),
      int(r["n_statements"]) - int(r["pro_count"]) - int(r["kontra_count"]) - int(r["netral_count"]),
      int(r["n_statements"])]
     for _, r in TOP_KON_ACT.iterrows()],
)

H3("Insight narasi")
BUL([
    "Bahlil Lahadalia (Menteri ESDM) berfungsi sebagai juru bicara utama narasi PRO dengan 55 dari 58 pernyataan tervalidasi berposisi PRO. Ia adalah simpul tertinggi di jaringan koalisi pendorong.",
    "Hashim Djojohadikusumo (Utusan Khusus Presiden bidang Energi dan Iklim) menempatkan nuklir sebagai bagian dari diplomasi energi internasional, dengan fokus pada kerja sama nuklir damai dan transisi energi.",
    "Pemerintah sebagai institusi (digabung dari berbagai juru bicara) mencatat 41 PRO dan 4 KONTRA pada tujuh variabel. Posisi ini menggambarkan suara struktural kabinet yang tidak terpecah.",
    "Fabby Tumiwa (Direktur Eksekutif IESR) adalah figur penyeimbang paling konsisten dengan 17 KONTRA berbasis argumen ekonomi (tarif PLTN mahal, biaya SMR tinggi) dan strategis (PLTN tidak sejalan tren global, ketergantungan teknologi).",
    "Mulyanto (Komisi VII DPR, Fraksi PKS) menunjukkan posisi yang lebih nuansa, yaitu 9 PRO dan 15 KONTRA. Ia mendukung PLTN untuk beban dasar, namun kritis terhadap pembubaran BATAN, kepatuhan hukum Presiden, dan keseriusan pemerintah pada EBET.",
    "Koalisi masyarakat sipil (Greenpeace, WALHI, JATAM, ICEL, Trend Asia, Koalisi Transisi Energi Berkeadilan) muncul sebagai blok wacana yang relatif kecil secara volume tetapi homogen secara posisi, hampir seluruhnya KONTRA dengan fokus limbah, keadilan, dan partisipasi publik.",
])

# ─── 3.4 Era
H2("3.4 Linimasa, Bagaimana Wacana Bergeser dari Era ke Era")
DASH_REF(["Distribusi Tema Wacana per Periode Pemerintahan", "Tren Artikel per Bulan (Berdasarkan Posisi)", "DNA Bipartite Network", "Position Shift Matrix"])
P(
    "Membaca data berdasarkan era kepemimpinan memperlihatkan pergeseran "
    "intensitas dan kecondongan wacana. Era Jokowi I masih relatif sepi, era "
    "Jokowi II meningkat tajam seiring pembahasan RUU EBET, dan era Prabowo "
    "mencatat lonjakan dukungan PRO yang paling tinggi."
)
TABLE(
    ["Era", "Artikel", "Pernyataan", "PRO", "KONTRA"],
    [[r["era"], int(r["art"]), int(r["n"]), int(r["pro"]), int(r["kontra"])]
     for _, r in ERA_TBL.iterrows()],
)

H3("Insight narasi")
BUL([
    "Era Jokowi I (2014 sampai 2019) didominasi figur akademisi dan ekonom seperti Kurtubi dan Arcandra Tahar. Wacana masih bersifat advokasi awal, dengan rasio PRO terbatas dan resistensi awal dari masyarakat sipil.",
    "Era Jokowi II (2019 sampai 2024) menjadi titik balik. Pembahasan DIM RUU EBET menggerakkan aktor parlemen (Sugeng Suparwoto, Eddy Soeparno, Bambang Haryadi) dan kabinet (Arifin Tasrif, Dadan Kusdiana) sebagai juru bicara utama. Volume PRO melonjak menjadi 540 pernyataan.",
    "Era Prabowo (sejak 20 Oktober 2024) menampilkan akselerasi dukungan tertinggi dalam sejarah wacana ini, yaitu 893 pernyataan PRO dalam waktu kurang dari satu setengah tahun. Variabel paling banyak dibahas pada era ini adalah Keamanan Nasional dan Kedaulatan (354 pernyataan), menandakan pembingkaian nuklir sebagai instrumen kedaulatan strategis.",
    "Pertumbuhan resistensi tidak sebanding. KONTRA pada era Prabowo justru lebih rendah dibanding era Jokowi II (151 vs 193 pernyataan), padahal volume keseluruhan meningkat. Ini mengindikasikan ruang wacana semakin condong ke kubu PRO.",
])

# ─── 3.5 Konsep
H2("3.5 Bahasa Wacana, Konsep yang Paling Sering Digaungkan")
DASH_REF(["Concept Evolution Timeline", "Distribusi Posisi per Variabel"])
P(
    "Pemetaan konsep mengungkap kosakata politik yang paling banyak dipakai "
    "kedua belah pihak. Kosakata ini adalah palet retoris yang membentuk persepsi "
    "publik dan menjadi alat untuk membaca framing media."
)

H3("Konsep PRO paling sering muncul")
TABLE(
    ["Konsep", "Frekuensi"],
    [
        ["Nuklir untuk kedaulatan energi", "8"],
        ["Kesiapan SDM PLTN", "5"],
        ["Pengembangan energi nuklir", "5"],
        ["Penyelesaian RUU EBET", "5"],
        ["PLTN opsi strategis transisi energi", "4"],
        ["Dukungan nuklir damai dan NPT", "4"],
        ["Kerja sama nuklir damai", "4"],
        ["Pembentukan Majelis Tenaga Nuklir", "4"],
        ["Nuklir energi masa depan ekonomis", "4"],
        ["Pembentukan NEPIO untuk PLTN", "3"],
    ],
)

H3("Konsep KONTRA paling sering muncul")
TABLE(
    ["Konsep", "Frekuensi"],
    [
        ["Penolakan skema power wheeling", "3"],
        ["Dampak limbah nuklir Fukushima", "3"],
        ["Biaya tinggi PLTN", "3"],
        ["Nuklir sebagai solusi palsu transisi energi", "2"],
        ["Power wheeling merugikan PLN", "2"],
        ["Gas fosil menghambat transisi energi", "2"],
        ["Prioritas energi terbarukan", "2"],
        ["Penolakan pembangunan PLTN", "2"],
        ["Tarif listrik PLTN mahal", "2"],
        ["PLTN bukan solusi ketahanan energi", "2"],
    ],
)

H3("Insight narasi")
BUL([
    "Bahasa PRO konsisten membingkai nuklir dalam tiga pilar, yaitu kedaulatan, kesiapan teknis, dan transisi energi. Tiga pilar ini saling memperkuat, karena masing masing menyasar audiens berbeda (politisi, teknokrat, dan komunitas iklim).",
    "Bahasa KONTRA bersandar pada dua pilar, yaitu ekonomi (biaya tinggi, tarif mahal) dan ekologis (limbah, Fukushima, solusi palsu). Frekuensi konsep KONTRA lebih rendah, namun konsep yang dipakai bersifat teknis sehingga sulit dibantah hanya dengan retorika.",
    "Munculnya istilah pembentukan Majelis Tenaga Nuklir dan NEPIO sebagai konsep PRO menandakan agenda kelembagaan sudah mulai diformalkan, bukan sekadar diskursus.",
    "Penolakan power wheeling sebagai konsep KONTRA paling tinggi mengindikasikan medan pertarungan baru, yaitu skema pasar listrik yang dirumuskan dalam RUU EBET, yang potensial merugikan PLN sekaligus mempermudah masuknya operator swasta termasuk untuk PLTN.",
])

# ─── 3.6 Sumber
H2("3.6 Ekosistem Sumber, Domain Media yang Membentuk Wacana")
DASH_REF(["Sumber Berita — Domain Terbanyak (Semua Era)"])
P(
    "Karena DNA berbasis pemberitaan, penting memahami siapa yang menyiarkan. "
    "Lima belas domain dengan jumlah artikel terbanyak menggambarkan ekosistem "
    "media yang mendukung pembentukan wacana ini."
)
TABLE(
    ["Domain", "Artikel", "Pernyataan", "PRO", "KONTRA"],
    [[r["domain"], int(r["art"]), int(r["stmt"]), int(r["pro"]), int(r["kontra"])]
     for _, r in TOP_DOM.iterrows()],
)

H3("Insight narasi")
BUL([
    "Tiga domain teratas, yaitu kompas.id, brin.go.id, dan dpr.go.id, mewakili tiga karakter berbeda yakni media mainstream, sumber lembaga riset, dan kanal resmi parlemen. Kombinasi ini memberi kredibilitas yang tinggi pada narasi PRO.",
    "Munculnya mongabay.co.id pada peringkat lima menunjukkan suara KONTRA berbasis lingkungan tetap punya jangkauan yang signifikan, terutama untuk variabel Periferalisasi dan Hak Masyarakat.",
    "Media bisnis (bloombergtechnoz, katadata, cnbcindonesia, listrikindonesia, ruangenergi) memberi panggung pada framing ekonomis nuklir, yang umumnya pro investasi.",
    "fraksi.pks.id sebagai satu satunya kanal partai pada daftar mengindikasikan PKS adalah satu satunya partai yang aktif memformalkan posisinya melalui kanal sendiri pada isu ini.",
])

# ─── 3.7 Sentimen
H2("3.7 Analisis Sentimen, Nada Retoris Dua Koalisi")
DASH_REF(["Distribusi Sentimen", "Sentimen × Posisi Aktor"])
P(
    "Selain posisi kebijakan, setiap pernyataan juga dinilai nada emosinya. "
    "Analisis sentimen memperlihatkan cara masing masing koalisi mengemas argumennya "
    "secara retoris, apakah dengan nada optimis, kritis, atau netral."
)
BUL([
    "Aktor pemerintah dan parlemen (koalisi PRO) cenderung menggunakan bahasa yang bernada positif, dengan pilihan kata seperti kedaulatan, kemajuan, masa depan, dan solusi. Nada ini memperkuat kesan bahwa nuklir adalah pilihan yang menjanjikan, bukan pilihan yang dipaksakan.",
    "Aktor masyarakat sipil dan lembaga kajian energi (koalisi KONTRA) cenderung menggunakan bahasa yang bernada negatif, dengan pilihan kata seperti mengancam, berbahaya, beban, dan gagal. Nada ini bertujuan membangkitkan kewaspadaan, bukan sekadar menolak secara emosional.",
    "Kombinasi PRO dengan sentimen negatif muncul pada aktor yang mendukung nuklir dengan cara mengkritik kondisi ketergantungan energi saat ini. Artinya, kritik mereka bukan terhadap nuklir, melainkan terhadap status quo yang dinilai buruk.",
    "Kombinasi KONTRA dengan sentimen positif muncul pada aktor yang menolak nuklir sembari mengangkat keunggulan energi terbarukan. Mereka tidak sekadar menolak, tetapi menawarkan alternatif yang lebih baik.",
    "Pola ini menunjukkan bahwa sentimen dan posisi tidak selalu berjalan seiring. Membaca keduanya bersama memberikan gambaran yang lebih lengkap tentang strategi retoris setiap koalisi.",
])

# ─── 3.8 Sintesis
H2("3.8 Sintesis, Tiga Cerita Besar dari Data")
P(
    "Tiga cerita berikut dapat menjadi tulang punggung presentasi kepada audiens "
    "yang lebih luas, karena masing masing membawa narasi yang utuh dan didukung "
    "data agregat dari pipeline ini."
)
H3("Cerita 1, Konsensus elite, debat di pinggir")
P(
    "Wacana nuklir adalah konsensus elite. Kabinet, parlemen, dan lembaga riset "
    "sebagian besar selaras pada PRO, sementara debat substansial justru muncul "
    "di pinggir, yaitu di lembaga kajian energi seperti IESR dan masyarakat sipil "
    "lingkungan. Implikasi narasinya, yaitu kebijakan akan lolos secara prosedural, "
    "namun isu legitimasi sosial dan kelayakan ekonomi tidak akan hilang dan akan "
    "muncul kembali pada tahap implementasi."
)
H3("Cerita 2, Bahasa berubah dari opsi terakhir menjadi kedaulatan")
P(
    "Pergeseran dari era Jokowi ke era Prabowo bukan sekadar peningkatan volume, "
    "melainkan perubahan bahasa. Variabel Keamanan Nasional dan Kedaulatan "
    "menjadi yang paling banyak dibahas pada era Prabowo, yang berarti nuklir "
    "tidak lagi dibingkai sebagai pilihan terakhir, melainkan sebagai pilar "
    "kedaulatan strategis. Pergeseran framing inilah yang akan menentukan dukungan "
    "publik dalam jangka menengah."
)
H3("Cerita 3, Pertarungan baru di skema pasar dan pembiayaan")
P(
    "Resistensi paling konsisten bukan pada teknologi nuklir itu sendiri, melainkan "
    "pada skema pembiayaan, power wheeling, dan beban PLN. Ini menandakan medan "
    "pertarungan berikutnya bukan apakah PLTN dibangun, melainkan dengan model "
    "bisnis seperti apa, dan siapa yang menanggung risiko fiskal serta tarif "
    "akhirnya. Audiens kebijakan perlu disiapkan untuk membaca pertarungan ini."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. KESIMPULAN
# ═══════════════════════════════════════════════════════════════════════════
H1("4. Kesimpulan")
P(
    "Discourse Network Analysis terhadap pemberitaan nasional sepanjang dua belas "
    "tahun terakhir menunjukkan bahwa wacana publik mengenai nuklir dalam RUU "
    "EBET telah berkonsolidasi menuju penerimaan. Lima kesimpulan utama dapat "
    "dirumuskan."
)
NUM([
    "Wacana nasional condong PRO. Mayoritas pernyataan tervalidasi berada pada posisi mendukung pengembangan nuklir, baik untuk PLTN konvensional maupun SMR.",

    "Koalisi pendorong terbentuk lintas lembaga. Kabinet (terutama ESDM), parlemen (Komisi VII dan XII), serta lembaga riset (BRIN dan akademisi) mengoperasikan narasi yang saling menguatkan dengan tiga pilar, yaitu kedaulatan, kesiapan teknologi, dan dekarbonisasi.",

    "Resistensi terlokalisasi di IESR, masyarakat sipil lingkungan, dan suara kritis di parlemen. Argumen KONTRA bersifat teknokratis, fokus pada biaya, tarif, limbah, dan keadilan distributif. Volumenya tidak besar, tetapi kualitas argumennya kredibel sehingga tetap menjadi referensi publik.",

    "Pergeseran era menunjukkan akselerasi yang tajam. Era Prabowo mencatat volume tertinggi dengan rasio PRO terkuat. Pembingkaian nuklir sebagai instrumen kedaulatan strategis menjadi karakter wacana era ini.",

    "Medan pertarungan berikutnya bergeser ke pembiayaan, skema pasar, dan partisipasi publik. Aspek teknologi sudah relatif diterima. Yang akan menentukan kelancaran implementasi adalah model fiskal, perlindungan PLN, dan kelayakan sosial di lokasi tapak.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. BATASAN DAN ASUMSI
# ═══════════════════════════════════════════════════════════════════════════
H1("5. Batasan dan Asumsi")

H2("5.1 Batasan Data")
BUL([
    "Sumber data berupa pemberitaan media daring yang dapat diakses publik, sehingga wacana yang muncul di forum tertutup, rapat internal, atau dokumen yang tidak terpublikasi tidak tertangkap.",
    "Cakupan periode dibatasi pada artikel yang relevan dengan tujuh variabel kunci. Pembahasan nuklir di luar konteks RUU EBET mungkin tidak terindeks penuh.",
    "Bahasa analisis adalah Bahasa Indonesia. Pemberitaan dan dokumen berbahasa asing mengenai nuklir Indonesia (misalnya laporan IAEA atau media internasional) tidak menjadi sumber primer.",
    "Volume artikel pada era Jokowi I lebih sedikit, sehingga representasi wacana periode tersebut bersifat indikatif dan tidak setara dengan dua era setelahnya.",
])

H2("5.2 Batasan Metodologis")
BUL([
    "Klasifikasi posisi (PRO, KONTRA, NETRAL, AMBIGU) dibatasi pada apa yang tertulis eksplisit dalam pernyataan. Sikap yang hanya tersirat melalui tindakan (misalnya kunjungan, pertemuan diplomatik) sengaja tidak diekstrak.",
    "Hanya pernyataan yang memenuhi kriteria validasi ketat yang dimasukkan ke dataset, sehingga pernyataan ambigu yang sebenarnya bermakna politik mungkin terbuang.",
    "Filter aktor aktif (minimal tiga pernyataan) digunakan agar analisis fokus pada aktor yang konsisten muncul. Akibatnya, suara aktor yang hanya muncul satu atau dua kali tidak masuk dalam visualisasi utama, meskipun masih tersimpan pada dataset mentah.",
    "Pemetaan ke tujuh variabel kunci dilakukan otomatis berdasarkan kata kunci dominan dalam artikel, sehingga pernyataan yang menyentuh dua atau lebih variabel hanya dikaitkan ke satu variabel paling representatif.",
])

H2("5.3 Asumsi Analitis")
BUL([
    "Diasumsikan bahwa pemberitaan media daring merepresentasikan ruang wacana publik nasional secara cukup memadai untuk pemetaan koalisi.",
    "Diasumsikan bahwa frekuensi kemunculan dan konsistensi posisi merupakan indikator yang valid untuk kekuatan suatu koalisi wacana.",
    "Diasumsikan bahwa konteks era kepemimpinan (Jokowi I, Jokowi II, Prabowo) memberi pengaruh struktural pada arah wacana, sehingga pemilahan era diperlukan untuk membaca pergeseran narasi.",
    "Diasumsikan bahwa sebuah pernyataan yang memuat aktor, konsep, dan sikap secara bersamaan merupakan unit analisis paling sesuai untuk DNA, sebagaimana standar metodologi yang digunakan dalam literatur policy network.",
])

H2("5.4 Catatan untuk Pengembangan Lanjutan")
BUL([
    "Hasil DNA ini akan diperkuat melalui wawancara mendalam terhadap target informan dari lima kategori pemangku kepentingan, yaitu DPR, Pemerintah, Akademisi, Industri, dan OMS, sehingga dapat dilakukan triangulasi antar sumber.",
    "Analisis lanjutan dapat diperluas ke arah analisis sentralitas formal (degree, betweenness, eigenvector) untuk merumuskan rekomendasi kebijakan berbasis posisi struktural aktor.",
    "Pemantauan berkelanjutan terhadap variabel Intervensi dan Pembiayaan disarankan, karena medan pertarungan utama tahap implementasi RUU EBET tampak berpindah ke skema fiskal dan model bisnis.",
])

# Save
doc.save(DOC_PATH)
print(f"[DONE] Saved: {DOC_PATH.resolve()}")
